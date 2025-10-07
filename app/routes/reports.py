"""报告生成路由"""
import os
from datetime import datetime
from flask import Blueprint, request, jsonify, current_app, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from app import db
from app.models.expense import Expense
from app.models.report import Report
from app.models.category import Category
from sqlalchemy import func
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端
import matplotlib.pyplot as plt
from matplotlib import font_manager
import io

bp = Blueprint('reports', __name__)

# 设置中文字体（如果有的话）
plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


@bp.route('/generate', methods=['POST'])
@jwt_required()
def generate_report():
    """生成分析报告"""
    current_user_id = get_jwt_identity()
    
    data = request.get_json()
    if not data:
        return jsonify({'error': '请提供报告参数'}), 400
    
    title = data.get('title', f'消费分析报告_{datetime.now().strftime("%Y%m%d")}')
    format_type = data.get('format', 'pdf')  # pdf, docx
    start_date = data.get('start_date')
    end_date = data.get('end_date')
    
    try:
        # 构建查询
        query = Expense.query.filter_by(user_id=current_user_id)
        
        if start_date:
            query = query.filter(Expense.time >= datetime.fromisoformat(start_date))
        if end_date:
            query = query.filter(Expense.time <= datetime.fromisoformat(end_date))
        
        expenses = query.all()
        
        if not expenses:
            return jsonify({'error': '没有数据可生成报告'}), 400
        
        # 生成图表
        charts = []
        
        # 1. 类别占比饼图
        category_data = db.session.query(
            Category.name,
            func.sum(Expense.amount).label('total')
        ).join(Expense).filter(Expense.user_id == current_user_id).group_by(Category.id).all()
        
        if category_data:
            fig, ax = plt.subplots(figsize=(8, 6))
            labels = [r.name for r in category_data]
            sizes = [r.total for r in category_data]
            ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
            ax.set_title('消费类别占比')
            
            chart_path = os.path.join(current_app.config['REPORT_FOLDER'], f'chart_pie_{current_user_id}_{datetime.now().timestamp()}.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()
            charts.append(chart_path)
        
        # 2. 消费趋势折线图
        daily_data = db.session.query(
            func.date(Expense.time).label('date'),
            func.sum(Expense.amount).label('total')
        ).filter_by(user_id=current_user_id).group_by('date').order_by('date').all()
        
        if daily_data:
            fig, ax = plt.subplots(figsize=(10, 6))
            dates = [r.date for r in daily_data]
            amounts = [r.total for r in daily_data]
            ax.plot(dates, amounts, marker='o')
            ax.set_title('消费趋势')
            ax.set_xlabel('日期')
            ax.set_ylabel('金额')
            plt.xticks(rotation=45)
            
            chart_path = os.path.join(current_app.config['REPORT_FOLDER'], f'chart_trend_{current_user_id}_{datetime.now().timestamp()}.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()
            charts.append(chart_path)
        
        # 生成报告文件
        if format_type == 'pdf':
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'report_{current_user_id}_{timestamp}.pdf'
            filepath = os.path.join(current_app.config['REPORT_FOLDER'], filename)
            
            c = canvas.Canvas(filepath, pagesize=A4)
            width, height = A4
            
            # 标题
            c.setFont("Helvetica-Bold", 20)
            c.drawString(50, height - 50, title)
            
            # 统计信息
            c.setFont("Helvetica", 12)
            y_position = height - 100
            
            total_amount = sum(e.amount for e in expenses)
            avg_amount = total_amount / len(expenses)
            
            c.drawString(50, y_position, f"Total Expenses: {len(expenses)}")
            y_position -= 20
            c.drawString(50, y_position, f"Total Amount: {total_amount:.2f}")
            y_position -= 20
            c.drawString(50, y_position, f"Average Amount: {avg_amount:.2f}")
            y_position -= 40
            
            # 插入图表
            for chart_path in charts:
                if y_position < 200:
                    c.showPage()
                    y_position = height - 50
                
                try:
                    img = ImageReader(chart_path)
                    c.drawImage(img, 50, y_position - 300, width=500, height=250, preserveAspectRatio=True)
                    y_position -= 320
                except Exception as e:
                    current_app.logger.error(f'插入图表失败: {e}')
            
            c.save()
            
        elif format_type == 'docx':
            from docx import Document
            from docx.shared import Inches
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'report_{current_user_id}_{timestamp}.docx'
            filepath = os.path.join(current_app.config['REPORT_FOLDER'], filename)
            
            doc = Document()
            doc.add_heading(title, 0)
            
            # 统计信息
            total_amount = sum(e.amount for e in expenses)
            avg_amount = total_amount / len(expenses)
            
            doc.add_paragraph(f'消费记录总数: {len(expenses)}')
            doc.add_paragraph(f'消费总金额: {total_amount:.2f}')
            doc.add_paragraph(f'平均消费金额: {avg_amount:.2f}')
            
            # 插入图表
            for chart_path in charts:
                try:
                    doc.add_picture(chart_path, width=Inches(6))
                except Exception as e:
                    current_app.logger.error(f'插入图表失败: {e}')
            
            doc.save(filepath)
        
        else:
            return jsonify({'error': '不支持的报告格式'}), 400
        
        # 保存报告记录
        report = Report(
            user_id=current_user_id,
            title=title,
            file_path=filepath,
            format=format_type
        )
        db.session.add(report)
        db.session.commit()
        
        current_app.logger.info(f'报告生成成功: {filename}')
        
        return jsonify({
            'message': '报告生成成功',
            'report': report.to_dict()
        }), 201
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'报告生成失败: {e}')
        return jsonify({'error': '报告生成失败', 'message': str(e)}), 500


@bp.route('/<int:report_id>/download', methods=['GET'])
@jwt_required()
def download_report(report_id):
    """下载报告"""
    current_user_id = get_jwt_identity()
    
    report = Report.query.filter_by(id=report_id, user_id=current_user_id).first()
    
    if not report:
        return jsonify({'error': '报告不存在'}), 404
    
    if not os.path.exists(report.file_path):
        return jsonify({'error': '报告文件不存在'}), 404
    
    return send_file(
        report.file_path,
        as_attachment=True,
        download_name=os.path.basename(report.file_path)
    )


@bp.route('/', methods=['GET'])
@jwt_required()
def get_reports():
    """获取报告列表"""
    current_user_id = get_jwt_identity()
    
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    
    pagination = Report.query.filter_by(user_id=current_user_id)\
        .order_by(Report.created_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'reports': [r.to_dict() for r in pagination.items],
        'total': pagination.total,
        'page': page,
        'per_page': per_page,
        'pages': pagination.pages
    }), 200

